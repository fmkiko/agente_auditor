#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera los entregables finales de una auditoría:
  - auditoria/{project_name}_{YYYYMMDD_HHMM}/auditoria-tecnica.xlsx
  - auditoria/{project_name}_{YYYYMMDD_HHMM}/informe-ejecutivo.md
  - auditoria/{project_name}_{YYYYMMDD_HHMM}/informe-ejecutivo.docx
  - .docs/Reporte_Tecnico.xlsx
  - .docs/Reporte_Ejecutivo.docx

El script recibe un JSON con metadatos de la auditoría (audit_id, project_path,
project_name, stack, resumen) y una lista de hallazgos estructurados.
"""
import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, Reference, PieChart


SEVERITY_ORDER = {"Crítica": 1, "Alta": 2, "Media": 3, "Baja": 4, "Info": 5}
SEVERITY_FILL = {
    "Crítica": PatternFill(start_color="C00000", end_color="C00000", fill_type="solid"),
    "Alta": PatternFill(start_color="FF6600", end_color="FF6600", fill_type="solid"),
    "Media": PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid"),
    "Baja": PatternFill(start_color="92D050", end_color="92D050", fill_type="solid"),
    "Info": PatternFill(start_color="00B0F0", end_color="00B0F0", fill_type="solid"),
}
SEVERITY_DOC_COLOR = {
    "Crítica": RGBColor(192, 0, 0),
    "Alta": RGBColor(255, 102, 0),
    "Media": RGBColor(255, 192, 0),
    "Baja": RGBColor(146, 208, 80),
    "Info": RGBColor(0, 176, 240),
}
ORDER_BY_SEVERITY = lambda s: SEVERITY_ORDER.get(s, 99)


def normalize_severity(value):
    value = str(value).strip().lower()
    mapping = {
        "critica": "Crítica", "crítica": "Crítica", "critical": "Crítica", "crit": "Crítica",
        "alta": "Alta", "high": "Alta", "grave": "Alta",
        "media": "Media", "medio": "Media", "medium": "Media",
        "baja": "Baja", "low": "Baja", "leve": "Baja",
    }
    return mapping.get(value, value.capitalize())


def ensure_default_finding_fields(finding):
    return {
        "audit_id": finding.get("audit_id", ""),
        "agent": finding.get("agent", ""),
        "category": finding.get("category", ""),
        "severity": normalize_severity(finding.get("severity", "Media")),
        "file": finding.get("file", ""),
        "line": finding.get("line", ""),
        "title": finding.get("title", ""),
        "description": finding.get("description", ""),
        "recommendation": finding.get("recommendation", ""),
    }


def write_xlsx(output_path, findings):
    wb = Workbook()
    ws = wb.active
    ws.title = "Hallazgos"

    headers = ["Audit ID", "Agente", "Categoría", "Severidad", "Archivo", "Línea", "Título", "Descripción", "Recomendación"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="44546A", end_color="44546A", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for finding in findings:
        row = [
            finding["audit_id"],
            finding["agent"],
            finding["category"],
            finding["severity"],
            finding["file"],
            finding["line"],
            finding["title"],
            finding["description"],
            finding["recommendation"],
        ]
        ws.append(row)
        sev = finding["severity"]
        fill = SEVERITY_FILL.get(sev, SEVERITY_FILL["Info"])
        ws.cell(row=ws.max_row, column=4).fill = fill

    for col_idx in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = min(50, max(12, len(headers[col_idx - 1]) + 4))
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    # Hoja Resumen
    ws2 = wb.create_sheet("Resumen")
    ws2.append(["Métrica", "Valor"])
    ws2.append(["Total de hallazgos", len(findings)])
    for cell in ws2[1]:
        cell.font = Font(bold=True)
    for sev in ["Crítica", "Alta", "Media", "Baja", "Info"]:
        count = sum(1 for f in findings if f["severity"] == sev)
        ws2.append([f"Hallazgos {sev}", count])

    # Agregados por categoría y agente
    ws2.append([])
    ws2.append(["Categoría", "Cantidad"])
    categories = {}
    agents = {}
    for f in findings:
        categories[f["category"]] = categories.get(f["category"], 0) + 1
        agents[f["agent"]] = agents.get(f["agent"], 0) + 1
    for cat, cnt in sorted(categories.items(), key=lambda x: -x[1]):
        ws2.append([cat, cnt])
    ws2.append([])
    ws2.append(["Agente", "Cantidad"])
    for ag, cnt in sorted(agents.items(), key=lambda x: -x[1]):
        ws2.append([ag, cnt])

    ws2.column_dimensions["A"].width = 40
    ws2.column_dimensions["B"].width = 15

    # Gráfico de barras por severidad
    chart = BarChart()
    chart.type = "col"
    chart.style = 10
    chart.title = "Hallazgos por severidad"
    chart.y_axis.title = "Cantidad"
    chart.x_axis.title = "Severidad"
    severities = ["Crítica", "Alta", "Media", "Baja", "Info"]
    counts = [sum(1 for f in findings if f["severity"] == s) for s in severities]
    data = [["Severidad", "Cantidad"]] + [[s, c] for s, c in zip(severities, counts)]
    ws2_tmp = wb.create_sheet("_severidad_chart_")
    for row in data:
        ws2_tmp.append(row)
    data_ref = Reference(ws2_tmp, min_col=2, min_row=1, max_row=len(data))
    cats = Reference(ws2_tmp, min_col=1, min_row=2, max_row=len(data))
    chart.add_data(data_ref, titles_from_data=True)
    chart.set_categories(cats)
    chart.shape = 4
    ws2.add_chart(chart, "D2")

    wb.save(output_path)


def create_risk_chart(output_image_path, findings):
    severities = ["Crítica", "Alta", "Media", "Baja", "Info"]
    counts = [sum(1 for f in findings if f["severity"] == s) for s in severities]
    colors = ["#C00000", "#FF6600", "#FFC000", "#92D050", "#00B0F0"]

    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.bar(severities, counts, color=colors)
    ax.set_title("Distribución de hallazgos por severidad")
    ax.set_ylabel("Cantidad")
    for bar, count in zip(bars, counts):
        height = bar.get_height()
        if count > 0:
            ax.annotate(str(count), xy=(bar.get_x() + bar.get_width()/2, height),
                        xytext=(0, 3), textcoords="offset points", ha="center", va="bottom")
    ax.set_ylim(0, max(counts + [1]) * 1.15)
    plt.tight_layout()
    plt.savefig(output_image_path, dpi=150)
    plt.close(fig)


def write_docx(output_path, md_path, findings, meta):
    doc = Document()
    title = doc.add_heading(f"Reporte Ejecutivo – {meta['project_name']}", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Audit ID: {meta['audit_id']}")
    doc.add_paragraph(f"Fecha: {meta.get('fecha', datetime.now().strftime('%Y-%m-%d %H:%M'))}")
    doc.add_paragraph(f"Ruta auditada: {meta['project_path']}")

    doc.add_heading("1. Identificación del proyecto", level=1)
    doc.add_paragraph(
        f"Se realizó una auditoría integral del proyecto {meta['project_name']} ubicado en "
        f"{meta['project_path']}. El objetivo fue identificar riesgos de seguridad, deuda técnica "
        "y oportunidades de mejora en el código fuente."
    )

    doc.add_heading("2. Alcance y stack tecnológico", level=1)
    doc.add_paragraph(meta.get("stack", "No se especificó el stack tecnológico."))

    doc.add_heading("3. Valoración general", level=1)
    total = len(findings)
    crit = sum(1 for f in findings if f["severity"] == "Crítica")
    high = sum(1 for f in findings if f["severity"] == "Alta")
    p = doc.add_paragraph()
    p.add_run(f"Se detectaron {total} hallazgos: ")
    for label, count, color in [
        ("Críticos", crit, SEVERITY_DOC_COLOR["Crítica"]),
        ("Altos", high, SEVERITY_DOC_COLOR["Alta"]),
        ("Medios", sum(1 for f in findings if f["severity"] == "Media"), SEVERITY_DOC_COLOR["Media"]),
        ("Bajos", sum(1 for f in findings if f["severity"] == "Baja"), SEVERITY_DOC_COLOR["Baja"]),
    ]:
        run = p.add_run(f"{count} {label}, ")
        run.font.color.rgb = color
        run.font.bold = True

    if crit or high:
        doc.add_paragraph(
            "Se recomienda no desplegar la aplicación en producción hasta mitigar los riesgos críticos y altos."
        )
    else:
        doc.add_paragraph("No se detectaron riesgos críticos. Se sugiere abordar los hallazgos medios y bajos en el próximo sprint.")

    doc.add_heading("4. Gráfico de riesgos", level=1)
    chart_path = str(output_path) + ".chart.png"
    create_risk_chart(chart_path, findings)
    doc.add_picture(chart_path, width=Inches(5.5))
    try:
        os.remove(chart_path)
    except OSError:
        pass

    doc.add_heading("5. Principales riesgos", level=1)
    top = sorted(findings, key=lambda f: (ORDER_BY_SEVERITY(f["severity"]), f["title"]))[:10]
    for i, f in enumerate(top, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"[{f['severity']}] {f['title']}")
        run.bold = True
        run.font.color.rgb = SEVERITY_DOC_COLOR.get(f["severity"], RGBColor(0, 0, 0))
        p.add_run(f" — {f['description'][:150]}{'...' if len(f['description']) > 150 else ''}")

    doc.add_heading("6. Recomendaciones prioritarias", level=1)
    recommendations = []
    for f in findings:
        if f["severity"] in ("Crítica", "Alta") and f.get("recommendation"):
            recommendations.append((f["severity"], f["recommendation"]))
    if recommendations:
        for sev, rec in recommendations[:8]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{sev}] ")
            run.bold = True
            run.font.color.rgb = SEVERITY_DOC_COLOR.get(sev, RGBColor(0, 0, 0))
            p.add_run(rec[:250])
    else:
        doc.add_paragraph("No se generaron recomendaciones prioritarias.")

    doc.add_heading("7. Próximos pasos", level=1)
    steps = [
        "Corregir hallazgos críticos y altos antes de cualquier despliegue.",
        "Actualizar dependencias con vulnerabilidades conocidas.",
        "Implementar autenticación, autorización y configuraciones seguras.",
        "Incorporar tests, linters y pipelines de CI/CD.",
        "Revisar y cerrar hallazgos medios en el siguiente ciclo de desarrollo.",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Anexo – Matriz resumida", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Severidad"
    hdr_cells[1].text = "Cantidad"
    hdr_cells[2].text = "% del total"
    hdr_cells[3].text = "Prioridad"
    for sev in ["Crítica", "Alta", "Media", "Baja", "Info"]:
        count = sum(1 for f in findings if f["severity"] == sev)
        row_cells = table.add_row().cells
        row_cells[0].text = sev
        row_cells[1].text = str(count)
        row_cells[2].text = f"{count/total*100:.1f}%" if total else "0%"
        row_cells[3].text = "Inmediata" if sev in ("Crítica", "Alta") else ("Corto plazo" if sev == "Media" else "Backlog")

    doc.save(output_path)

    # Guardar markdown equivalente
    md_lines = [
        f"# Reporte Ejecutivo – {meta['project_name']}",
        "",
        f"**Audit ID:** {meta['audit_id']}",
        f"**Fecha:** {meta.get('fecha', datetime.now().strftime('%Y-%m-%d %H:%M'))}",
        f"**Ruta auditada:** {meta['project_path']}",
        "",
        "## 1. Identificación del proyecto",
        f"Auditoría integral del proyecto {meta['project_name']}. Objetivo: identificar riesgos, deuda técnica y oportunidades de mejora.",
        "",
        "## 2. Alcance y stack tecnológico",
        meta.get("stack", "No especificado."),
        "",
        "## 3. Valoración general",
        f"Se detectaron **{total}** hallazgos: {crit} críticos, {high} altos, {sum(1 for f in findings if f['severity']=='Media')} medios, {sum(1 for f in findings if f['severity']=='Baja')} bajos.",
        "",
        "## 4. Principales riesgos",
    ]
    for f in top:
        md_lines.append(f"- **[{f['severity']}] {f['title']}** — {f['description'][:200]}{'...' if len(f['description'])>200 else ''}")
    md_lines += [
        "",
        "## 5. Recomendaciones prioritarias",
    ]
    for sev, rec in recommendations[:10]:
        md_lines.append(f"- **[{sev}]** {rec}")
    md_lines += ["", "## 6. Próximos pasos"]
    for s in steps:
        md_lines.append(f"- {s}")
    Path(md_path).write_text("\n".join(md_lines), encoding="utf-8")


def main():
    parser = argparse.ArgumentParser(description="Genera entregables de auditoría")
    parser.add_argument("--input", "-i", required=True, help="JSON con metadatos y hallazgos")
    parser.add_argument("--output-dir", "-o", required=True, help="Carpeta de salida")
    parser.add_argument("--docs-dir", "-d", default=".docs", help="Carpeta .docs para copias finales")
    args = parser.parse_args()

    payload = json.loads(Path(args.input).read_text(encoding="utf-8"))
    meta = payload.get("meta", {})
    raw_findings = payload.get("findings", [])
    findings = [ensure_default_finding_fields(f) for f in raw_findings]

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    docs_dir = Path(args.docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    xlsx_path = out_dir / "auditoria-tecnica.xlsx"
    md_path = out_dir / "informe-ejecutivo.md"
    docx_path = out_dir / "informe-ejecutivo.docx"

    write_xlsx(str(xlsx_path), findings)
    write_docx(str(docx_path), str(md_path), findings, meta)

    # Copias con nombres solicitados
    report_xlsx = docs_dir / "Reporte_Tecnico.xlsx"
    report_docx = docs_dir / "Reporte_Ejecutivo.docx"
    import shutil
    shutil.copy2(str(xlsx_path), str(report_xlsx))
    shutil.copy2(str(docx_path), str(report_docx))

    outputs = {
        "auditoria_dir": str(out_dir.resolve()),
        "xlsx": str(xlsx_path.resolve()),
        "docx": str(docx_path.resolve()),
        "md": str(md_path.resolve()),
        "report_xlsx": str(report_xlsx.resolve()),
        "report_docx": str(report_docx.resolve()),
    }
    print(json.dumps(outputs, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
